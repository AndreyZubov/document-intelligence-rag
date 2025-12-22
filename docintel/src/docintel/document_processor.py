"""Document processing and text extraction."""

import logging
from pathlib import Path
from typing import BinaryIO
from uuid import UUID

import pypdf
import docx
import markdown

from docintel.config import settings
from docintel.models import Document, DocumentChunk

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document parsing and chunking."""

    def __init__(self, chunk_size: int = settings.chunk_size, chunk_overlap: int = settings.chunk_overlap):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def extract_text(self, file: BinaryIO, filename: str) -> str:
        """Extract text from various document formats."""
        suffix = Path(filename).suffix.lower()

        try:
            if suffix == ".pdf":
                return self._extract_pdf(file)
            elif suffix in [".docx", ".doc"]:
                return self._extract_docx(file)
            elif suffix in [".md", ".markdown"]:
                return self._extract_markdown(file)
            elif suffix == ".txt":
                return file.read().decode("utf-8", errors="ignore")
            else:
                raise ValueError(f"Unsupported file format: {suffix}")
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            raise

    def _extract_pdf(self, file: BinaryIO) -> str:
        """Extract text from PDF file."""
        reader = pypdf.PdfReader(file)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)

    def _extract_docx(self, file: BinaryIO) -> str:
        """Extract text from DOCX file."""
        doc = docx.Document(file)
        return "\n\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])

    def _extract_markdown(self, file: BinaryIO) -> str:
        """Extract text from Markdown file."""
        content = file.read().decode("utf-8", errors="ignore")
        # Convert to HTML then strip tags for plain text, or keep markdown as-is
        return content

    def create_chunks(self, text: str, document_id: UUID) -> list[DocumentChunk]:
        """Split text into overlapping chunks."""
        if not text or not text.strip():
            logger.warning(f"Empty text provided for document {document_id}")
            return []

        # Simple character-based chunking with overlap
        chunks = []
        text = text.strip()
        start = 0
        chunk_index = 0

        while start < len(text):
            end = start + self.chunk_size
            chunk_text = text[start:end]

            # Try to break at sentence or word boundary if possible
            if end < len(text):
                # Look for sentence end
                last_period = chunk_text.rfind(". ")
                last_newline = chunk_text.rfind("\n")
                last_space = chunk_text.rfind(" ")

                break_point = max(last_period, last_newline, last_space)
                if break_point > self.chunk_size * 0.5:  # Only break if we're past halfway
                    chunk_text = chunk_text[: break_point + 1].strip()
                    end = start + len(chunk_text)

            if chunk_text:
                chunk = DocumentChunk(
                    id=f"{document_id}_{chunk_index}",
                    document_id=document_id,
                    content=chunk_text.strip(),
                    chunk_index=chunk_index,
                    metadata={"start_pos": start, "end_pos": end},
                )
                chunks.append(chunk)
                chunk_index += 1

            # Move start position with overlap
            start = end - self.chunk_overlap

        logger.info(f"Created {len(chunks)} chunks for document {document_id}")
        return chunks

    def process_document(self, file: BinaryIO, filename: str) -> tuple[Document, list[DocumentChunk]]:
        """Process a document end-to-end: extract text and create chunks."""
        text = self.extract_text(file, filename)

        if not text or len(text.strip()) < 10:
            raise ValueError(f"Document {filename} appears to be empty or too short")

        document = Document(
            filename=filename,
            content=text[:1000],  # Store preview only
            metadata={"original_length": len(text), "format": Path(filename).suffix},
        )

        chunks = self.create_chunks(text, document.id)
        document.chunk_count = len(chunks)

        logger.info(f"Processed document {filename}: {len(chunks)} chunks created")
        return document, chunks
